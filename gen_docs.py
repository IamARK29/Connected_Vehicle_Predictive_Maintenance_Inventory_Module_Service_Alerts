"""Generate professional Word documents for AutoPredict documentation suite."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOCS  = Path(r"C:\Users\Developer\Projects\AutoPredict\docs")
DIAG  = DOCS / "diagrams"
DOCS.mkdir(parents=True, exist_ok=True)

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0F, 0x20, 0x44)
BLUE   = RGBColor(0x1E, 0x57, 0x99)
CYAN   = RGBColor(0x00, 0xB4, 0xD8)
GREEN  = RGBColor(0x14, 0xA4, 0x4D)
ORG    = RGBColor(0xE7, 0x6F, 0x51)
RED    = RGBColor(0xE6, 0x39, 0x46)
GRAY   = RGBColor(0x64, 0x74, 0x8B)
LGRAY  = RGBColor(0xE2, 0xE8, 0xF0)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1A, 0x20, 0x2C)

# ── Helpers ───────────────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    return doc

def set_para_color(para, rgb):
    for run in para.runs:
        run.font.color.rgb = rgb

def heading(doc, text, level=1, color=NAVY, size=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    return p

def body(doc, text, color=DARK, size=10.5):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph('')

def caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = GRAY

def add_image(doc, filename, width=Inches(6.5), caption_text=None):
    img_path = DIAG / filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
        if caption_text:
            caption(doc, caption_text)
    else:
        body(doc, f'[Diagram not found: {filename}]', color=RED)

def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None, header_color='0F2044'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    shade_row(hdr_row, header_color)
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_data in rows:
        row = table.add_row()
        for i, val in enumerate(r_data):
            cell = row.cells[i]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    return table

def cover_page(doc, title, subtitle, version='v1.0', date='July 2025',
               project='AutoPredict — Predictive Maintenance Platform'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MG MOTOR INDIA')
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.font.bold = True
    run.font.character_spacing = 200

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project)
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE
    run.font.bold = True

    spacer(doc, 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.font.bold = True

    spacer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY
    run.font.italic = True

    spacer(doc, 4)

    meta = [('Version', version), ('Date', date),
            ('Status', 'Confidential — Internal Use'), ('Platform', 'AutoPredict v2.0')]
    for k, v in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{k}:  ')
        r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
        r2 = p.add_run(v)
        r2.font.size = Pt(10); r2.font.color.rgb = GRAY

    doc.add_page_break()

# ── Doc 1 — Technical Architecture ───────────────────────────────────────────
def doc_tech_arch():
    doc = new_doc()
    cover_page(doc, 'Technical Architecture Document',
               'System design, technology stack, and architectural decisions')

    heading(doc, '1. Executive Summary')
    body(doc, 'AutoPredict is a full-stack predictive maintenance platform built for MG Motor India. '
         'It processes vehicle telemetry from TBox ECUs in real time, applies LightGBM machine learning '
         'models to predict component failures up to 30 days in advance, and surfaces actionable '
         'maintenance alerts through a React-based dealer portal. The platform is designed for '
         'offline-first operation using CSV data stores, enabling deployment in environments with '
         'limited connectivity.')

    heading(doc, '2. System Architecture Overview', level=1)
    body(doc, 'The system is organised into five distinct layers, each with a clear responsibility boundary:')
    add_image(doc, 'diag_01_system_arch.png', width=Inches(6.5),
              caption_text='Figure 1 — Five-layer system architecture of AutoPredict')

    for layer, desc in [
        ('Layer 1 — Vehicle / Edge', 'MG vehicles equipped with TBox ECUs emit 200+ telemetry signals at 5 Hz via MQTT. Signals follow the MG TBox Big Data Specification with EV/PHEV/ICE schema variants.'),
        ('Layer 2 — Messaging & Storage', 'Eclipse Mosquitto broker receives MQTT telemetry. The platform uses pandas CSV files as the primary data store — a deliberate offline-first decision enabling operation without a live database.'),
        ('Layer 3 — ML & Business Logic', '8 LightGBM models provide RUL (Remaining Useful Life) predictions for major components. A Champion-Challenger registry with PSI drift monitoring governs model promotion. 4 physics-based EV engines (DCDC, motor, charging, thermal) complement the statistical models.'),
        ('Layer 4 — API Gateway', 'FastAPI with Uvicorn provides 36 REST endpoints across 6 routers. Authentication uses JWT HS256 with three-role RBAC (DEALER/OEM/ADMIN). slowapi enforces rate limits on auth endpoints.'),
        ('Layer 5 — Presentation', 'React 18 + TypeScript dealer portal with role-gated routing. TanStack Query v5 manages server state. Recharts renders analytics charts.'),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{layer}: ')
        r1.font.bold = True; r1.font.color.rgb = NAVY; r1.font.size = Pt(10.5)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)

    heading(doc, '3. Data Flow Pipeline', level=1)
    body(doc, 'Raw telemetry signals flow from TBox ECUs through the full stack to generate dealer-facing alerts:')
    add_image(doc, 'diag_02_data_flow.png', width=Inches(6.5),
              caption_text='Figure 2 — Data pipeline from TBox signal to dealer alert')
    body(doc, 'Note: The first CSV read on a cold start takes 55–65 seconds due to pandas '
         'loading and aggregating the full telemetry history. Subsequent reads are cached in memory.')

    heading(doc, '4. Authentication & RBAC Architecture', level=1)
    add_image(doc, 'diag_03_auth_rbac.png', width=Inches(6.5),
              caption_text='Figure 3 — JWT authentication flow and role hierarchy')
    body(doc, 'The JWT payload carries three claims critical to authorisation: sub (username), '
         'role (dealer/oem/admin), and dealer_code. Every database query filters on dealer_code, '
         'ensuring that a dealer with a valid JWT cannot access another dealer\'s data even if '
         'they intercept or forge a token with a different dealer_code claim.')

    heading(doc, '5. Technology Stack', level=1)
    add_table(doc,
        ['Layer', 'Technology', 'Version', 'Rationale'],
        [
            ('Frontend', 'React + TypeScript', '18 / 5', 'Component model fits the dashboard/portal pattern; TypeScript prevents runtime type errors across 11 pages'),
            ('Frontend', 'Vite', '5', 'Sub-second HMR; native ES modules; faster than CRA/Webpack for this project size'),
            ('Frontend', 'TanStack Query', 'v5', 'Automatic caching, stale-while-revalidate, and background refetch — eliminates manual useEffect data fetching'),
            ('Frontend', 'Recharts', '2', 'Pure SVG; passes props through to path elements enabling aria-label injection for WCAG compliance'),
            ('Frontend', 'Tailwind CSS', '3', 'Utility-first CSS eliminates naming overhead and keeps bundle size low via PurgeCSS'),
            ('Backend', 'FastAPI', '0.110', 'Auto-generates OpenAPI schema; native async; Pydantic validation at system boundary'),
            ('Backend', 'Uvicorn', '0.29', 'ASGI server; --reload during dev; single worker in prod matches CSV single-threaded constraint'),
            ('Backend', 'python-jose', '3.3', 'JWT HS256 encode/decode; battle-tested; integrates cleanly with FastAPI Depends pattern'),
            ('Backend', 'pandas', '2.x', 'Offline-first CSV engine; vectorised aggregation over millions of telemetry rows'),
            ('ML', 'LightGBM', '4.x', 'Gradient boosted trees; fastest training on tabular data; handles missing values natively'),
            ('ML', 'scikit-learn', '1.4', 'Feature pipeline, scaler, and model evaluation metrics'),
            ('ML', 'paho-mqtt', '1.6', 'MQTT client for TBox telemetry subscription'),
            ('Testing', 'Playwright', '1.61', 'Native TypeScript; built-in APIRequestContext; reliable cross-browser automation'),
            ('Testing', 'axe-core', '4.x', 'WCAG 2.1 AA automated accessibility scanning in-browser'),
        ],
        col_widths=[Inches(1.0), Inches(1.4), Inches(0.8), Inches(3.2)])

    heading(doc, '6. Deployment Architecture', level=1)
    add_image(doc, 'diag_04_deployment.png', width=Inches(6.5),
              caption_text='Figure 4 — Docker Compose deployment with three containers')
    add_table(doc,
        ['Service', 'Port', 'Technology', 'Purpose'],
        [
            ('frontend', '3000', 'React + Vite / Nginx', 'Serves compiled React SPA'),
            ('backend', '8001', 'FastAPI + Uvicorn', 'REST API, JWT auth, ML inference'),
            ('mqtt-broker', '1883', 'Eclipse Mosquitto', 'Vehicle telemetry ingestion'),
        ])

    heading(doc, '7. ML Pipeline', level=1)
    add_image(doc, 'diag_05_ml_pipeline.png', width=Inches(6.5),
              caption_text='Figure 5 — ML training and serving pipeline with Champion-Challenger registry')
    body(doc, 'The Champion-Challenger framework ensures safe model promotion: a new model (Challenger) '
         'runs in shadow mode against live traffic. PSI (Population Stability Index) monitors feature '
         'distribution drift. When PSI exceeds the configured threshold, the OEM analyst is notified '
         'via the Retrain Control UI and can promote the challenger to champion.')

    heading(doc, '8. API Architecture', level=1)
    add_image(doc, 'diag_06_api_routes.png', width=Inches(6.5),
              caption_text='Figure 6 — API router map showing 36 endpoints across 6 routers')

    heading(doc, '9. Architectural Decisions & Rationale', level=1)
    decisions = [
        ('CSV over Database', 'Enables fully offline deployment in dealerships with intermittent connectivity. pandas provides sufficient performance for the current fleet size (<50,000 vehicles). Migration path to PostgreSQL is documented in the Phase 4 roadmap.'),
        ('JWT over Session Cookies', 'Stateless auth simplifies horizontal scaling and eliminates server-side session storage. All auth state lives in localStorage — acceptable for the internal dealer portal threat model.'),
        ('LightGBM over Neural Networks', 'Tabular vehicle telemetry data with hand-crafted features performs better with gradient boosted trees than neural networks. LightGBM trains in minutes on commodity hardware; neural nets would require GPU resources not available at dealerships.'),
        ('workers=1 in Production', 'The CSV data store is not thread-safe for concurrent writes. Single-worker Uvicorn is a deliberate constraint, not a limitation. The Phase 4 database migration removes this constraint.'),
        ('React Query over Redux', 'Server state (API data) is separate from client state (UI). React Query handles the former with caching and background refetch; useState/Context handle the latter. Redux would add unnecessary boilerplate.'),
    ]
    for decision, rationale in decisions:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{decision}: ')
        r1.font.bold = True; r1.font.color.rgb = NAVY; r1.font.size = Pt(10.5)
        r2 = p.add_run(rationale)
        r2.font.size = Pt(10.5)

    doc.save(DOCS / '01_Technical_Architecture.docx')
    print('  ✓ 01_Technical_Architecture.docx')


# ── Doc 2 — Product Specification ────────────────────────────────────────────
def doc_product_spec():
    doc = new_doc()
    cover_page(doc, 'Product Specification Document',
               'Feature definitions, user roles, and functional scope')

    heading(doc, '1. Product Overview')
    body(doc, 'AutoPredict is a cloud-ready, offline-capable predictive maintenance platform '
         'purpose-built for MG Motor India. It connects vehicle telemetry from the TBox ECU '
         'network with machine learning models to predict component failures before they occur, '
         'enabling proactive service scheduling at the dealer level and fleet-wide analytics '
         'at the OEM level.')
    body(doc, 'The platform replaces reactive maintenance (fix after failure) with predictive '
         'maintenance (fix before failure), targeting a 30% reduction in unplanned breakdowns '
         'and a 20% improvement in service bay utilisation.')

    heading(doc, '2. User Roles & Permissions')
    body(doc, 'AutoPredict supports three distinct user roles with hierarchical access:')
    add_table(doc,
        ['Role', 'Scope', 'Key Capabilities', 'Restriction'],
        [
            ('DEALER', 'Own fleet (dealer_code scoped)', 'Dashboard, Alerts, Service Bay, Inventory', 'Cannot see other dealers\' data'),
            ('OEM', 'All dealers and all vehicles', 'Fleet Analytics, Model Health, EDA, What-If, Retrain', 'Cannot manage users'),
            ('ADMIN', 'Full system', 'User management, all OEM capabilities', 'No additional restriction'),
        ])

    heading(doc, '3. System Architecture', level=1)
    add_image(doc, 'diag_01_system_arch.png', width=Inches(6.5),
              caption_text='Figure 1 — AutoPredict five-layer system architecture')

    heading(doc, '4. Feature Specifications', level=1)

    heading(doc, '4.1 Authentication & Role Management', level=2)
    body(doc, 'JWT HS256-based authentication with 3-tier role hierarchy. Tokens carry sub, role, '
         'dealer_code, iat, and exp claims. HTTPBearer with auto_error=False returns HTTP 401 '
         'for missing credentials and HTTP 403 for valid token with insufficient role.')
    bullet(doc, 'Login page with username/password form and branded MG Motor split layout')
    bullet(doc, 'localStorage-based auth persistence (ap_token, ap_role, ap_dealer_code, ap_user)')
    bullet(doc, 'PrivateRoute and OemRoute guards in React Router v6')
    bullet(doc, 'slowapi rate limiting on /api/auth/token (5 requests/minute)')
    bullet(doc, 'Automatic redirect to /login on token expiry or 401 response')

    heading(doc, '4.2 Fleet Dashboard (Dealer)', level=2)
    body(doc, 'The primary dealer-facing page providing a fleet health overview with ML-driven '
         'risk scoring for every vehicle in the dealer\'s fleet.')
    bullet(doc, '4 KPI cards: Total Vehicles, Active Alerts, Due for Service, Fleet Health Score')
    bullet(doc, 'Sortable vehicle table with columns: VIN, Model, Year, Mileage, Last Service, Risk Score, Status')
    bullet(doc, 'Search/filter by VIN or model name')
    bullet(doc, 'Colour-coded risk status: Critical (red), High (orange), Medium (yellow), Good (green)')
    bullet(doc, 'Upcoming Service panel with next 7 days appointments')
    bullet(doc, 'Real-time data refresh every 5 minutes via TanStack Query')

    heading(doc, '4.3 Predictive Alerts', level=2)
    body(doc, 'Rule-based and ML-derived alerts surfaced in a filterable table, enabling dealers '
         'to prioritise vehicle maintenance by severity and component type.')
    bullet(doc, 'Severity filter: All / Critical / High / Medium / Low')
    bullet(doc, 'Time-window filter: Last 24h / 7 days / 30 days / 90 days')
    bullet(doc, 'Alert columns: VIN, Component, Predicted Failure Window, Confidence, Days Remaining')
    bullet(doc, 'Manual refresh button with debounce')
    bullet(doc, 'Empty state when no alerts match filters')

    heading(doc, '4.4 ML Prediction Models', level=2)
    body(doc, 'Eight LightGBM models provide component-specific Remaining Useful Life (RUL) '
         'predictions. Four additional physics-based engines handle EV-specific subsystems.')
    add_table(doc,
        ['Model', 'Target Component', 'Key Features Used', 'Output'],
        [
            ('brake_pad', 'Brake pads', 'Brake applications/km, decel G-force, pad wear signal', 'RUL in km'),
            ('engine_oil', 'Engine oil', 'Oil temp, viscosity proxy, km since last change', 'RUL in km'),
            ('tyre_wear', 'Tyres', 'Mileage, inflation pressure, cornering G', 'RUL in km'),
            ('battery_12v', '12V auxiliary battery', 'Voltage, temperature, charge cycles', 'RUL in days'),
            ('hv_battery_soh', 'HV battery (EV/PHEV)', 'SOC, charge cycles, thermal history', 'SoH %'),
            ('transmission', 'Transmission', 'Gear changes, temperature, fluid level proxy', 'RUL in km'),
            ('cooling_system', 'Cooling system', 'Coolant temp, fan duty cycle, flow rate', 'RUL in days'),
            ('driver_score', 'Driver behaviour', 'Harsh braking, rapid acceleration, speed percentiles', 'Score 0–100'),
            ('DCDC Converter', 'EV DCDC system', 'Input/output voltage, current, efficiency', 'Health %'),
            ('Motor Health', 'EV drive motor', 'Phase currents, temperatures, torque efficiency', 'Health %'),
            ('Charging Health', 'Charging system', 'Charge rate, thermal, BMS logs', 'Health %'),
            ('Thermal Management', 'Battery thermal', 'Cell temperatures, delta T, coolant flow', 'Health %'),
        ],
        col_widths=[Inches(1.3), Inches(1.4), Inches(2.2), Inches(1.5)])

    heading(doc, '4.5 Service Bay Management', level=2)
    bullet(doc, 'Visual bay grid showing bay number, current status (Free/Occupied/Reserved)')
    bullet(doc, 'Booking modal: VIN selection, service type, estimated duration, technician')
    bullet(doc, 'Appointment list with upcoming bookings sortable by date and urgency')
    bullet(doc, 'Integration with alert system — book directly from an alert record')

    heading(doc, '4.6 Inventory Management', level=2)
    bullet(doc, '7-tab inventory interface: Overview, Stock Levels, Reorder Plan, Demand Forecast, Slow Moving, EV Parts, Supplier Contacts')
    bullet(doc, 'Stock levels per part with min/max thresholds and current quantity')
    bullet(doc, 'ML-driven demand forecast with configurable horizon (30/60/90 days)')
    bullet(doc, 'Automatic reorder plan generated from demand forecast vs current stock')
    bullet(doc, 'EV-specific parts tab with battery pack and charging component inventory')

    heading(doc, '4.7 OEM Fleet Intelligence', level=2)
    add_image(doc, 'diag_06_api_routes.png', width=Inches(6.3),
              caption_text='Figure 2 — OEM API router structure with role access')
    bullet(doc, 'Fleet-wide overview with group-by: Model / Region / Year / Fuel Type')
    bullet(doc, 'Model Health heatmap: AUROC, PSI drift, Precision, Recall per model per region')
    bullet(doc, 'EDA dashboard: Feature distributions, correlation heatmap, outlier detection')
    bullet(doc, 'What-If analysis: Adjust feature values and observe predicted outcome change')
    bullet(doc, 'Demand breakdown: Per-part pivot table segmented by region and dealer code')

    heading(doc, '4.8 Model Retrain Control (OEM)', level=2)
    bullet(doc, 'Multi-select checkboxes for 8 LightGBM models + 4 EV physics engines')
    bullet(doc, 'Select All / Clear All controls for batch operations')
    bullet(doc, 'Training notes text field for documenting the rationale for retraining')
    bullet(doc, 'Submit triggers background training job with progress tracking')
    bullet(doc, 'Retrain history table: Model, timestamp, triggered by, status, AUROC before/after')

    heading(doc, '5. Non-Functional Requirements', level=1)
    add_table(doc,
        ['Requirement', 'Target', 'Notes'],
        [
            ('Dashboard load time', '< 3s (warm)', 'Cold CSV read 55–65s; cached thereafter'),
            ('Alert prediction latency', '< 500ms', 'In-memory model inference after warm-up'),
            ('Concurrent users', '50 per dealer', 'Single Uvicorn worker; Phase 4 scales to multi-worker'),
            ('Data retention', '24 months rolling', 'CSV files rotated monthly'),
            ('Uptime', '99.5%', 'Single-server deployment; Phase 4 adds HA'),
            ('WCAG Compliance', '2.1 AA', 'axe-core verified; all 8 pages pass'),
            ('Security', 'OWASP Top 10 awareness', '5 vulnerabilities documented; remediation in Phase 4'),
        ])

    doc.save(DOCS / '02_Product_Specification.docx')
    print('  ✓ 02_Product_Specification.docx')


# ── Doc 3 — PRD ───────────────────────────────────────────────────────────────
def doc_prd():
    doc = new_doc()
    cover_page(doc, 'Product Requirements Document',
               'Functional requirements, personas, MoSCoW prioritisation, and acceptance criteria')

    heading(doc, '1. Background & Problem Statement')
    body(doc, 'MG Motor India operates a growing fleet managed through a network of authorised '
         'dealers. Vehicle maintenance today is largely reactive — service happens after warning '
         'lights appear or customers report problems. This creates three business problems:')
    bullet(doc, 'Unplanned breakdowns erode customer trust and generate warranty costs')
    bullet(doc, 'Service bays are underutilised due to unpredictable demand spikes')
    bullet(doc, 'OEM teams have no visibility into fleet health trends across dealer networks')

    heading(doc, '2. Product Vision')
    body(doc, '"Enable every MG dealer to act on a vehicle failure 30 days before it happens — '
         'turning reactive maintenance into a proactive, data-driven service operation."')

    heading(doc, '3. User Personas', level=1)
    add_table(doc,
        ['Persona', 'Role', 'Goal', 'Pain Point'],
        [
            ('Rajesh Kumar', 'Dealer Service Manager\nMumbai dealership', 'Know which vehicles need service before customers complain', 'Currently relies on customer calls and warning lights'),
            ('Priya Sharma', 'OEM Fleet Data Engineer\nMG Motor India HQ', 'Monitor ML model quality and retrain when data drifts', 'No centralised view of model performance across fleet'),
            ('Arjun Mehta', 'System Administrator', 'Manage user accounts and access rights', 'Manual user management with no audit trail'),
        ])

    heading(doc, '4. Requirements', level=1)
    heading(doc, '4.1 MoSCoW Prioritisation', level=2)

    moscow = {
        'Must Have (P1)': [
            'JWT authentication with role-based access control',
            'Fleet dashboard with vehicle risk scores',
            'Predictive alerts with severity levels',
            'LightGBM ML models for 8 components',
            'CSV offline-first data storage',
            'OEM fleet analytics overview',
            'WCAG 2.1 AA accessibility compliance',
        ],
        'Should Have (P2)': [
            'Service bay booking management',
            'Inventory management with demand forecast',
            'OEM model health monitoring (AUROC, PSI)',
            'OEM What-If analysis tool',
            'Model retrain control with history',
            'EV/PHEV-specific health engines',
            'Champion-Challenger model serving',
        ],
        'Could Have (P3)': [
            'Mobile PWA for dealer technicians',
            'Real-time WebSocket push notifications',
            'PDF report export for service records',
            'Integration with DMS (Dealer Management System)',
            'Multi-language support (Hindi, regional)',
        ],
        "Won't Have (this release)": [
            'Customer-facing mobile app',
            'OTA vehicle software updates',
            'Insurance integration',
            'Carbon footprint tracking',
        ],
    }

    for category, items in moscow.items():
        p = doc.add_paragraph()
        r = p.add_run(category)
        r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(11)
        for item in items:
            bullet(doc, item)

    heading(doc, '5. Functional Requirements', level=1)
    reqs = [
        ('FR-001', 'Auth', 'Must', 'System shall authenticate users via username/password and return JWT'),
        ('FR-002', 'Auth', 'Must', 'System shall enforce role-based access: DEALER, OEM, ADMIN'),
        ('FR-003', 'Auth', 'Must', 'System shall scope all dealer queries to their dealer_code'),
        ('FR-004', 'Fleet', 'Must', 'System shall display fleet KPIs: total vehicles, active alerts, due for service'),
        ('FR-005', 'Fleet', 'Must', 'System shall show vehicle-level risk scores from ML models'),
        ('FR-006', 'Alerts', 'Must', 'System shall generate predictive alerts 30 days before predicted failure'),
        ('FR-007', 'Alerts', 'Must', 'System shall support severity filter: Critical/High/Medium/Low'),
        ('FR-008', 'ML', 'Must', 'System shall serve predictions from 8 LightGBM maintenance models'),
        ('FR-009', 'ML', 'Should', 'System shall monitor PSI drift and alert OEM when threshold exceeded'),
        ('FR-010', 'OEM', 'Should', 'System shall provide fleet-wide analytics grouped by model/region/year'),
        ('FR-011', 'OEM', 'Should', 'System shall allow OEM to trigger model retraining via UI'),
        ('FR-012', 'Inventory', 'Should', 'System shall forecast part demand using ML models'),
        ('FR-013', 'Admin', 'Should', 'System shall allow admin to create, list, and deactivate user accounts'),
        ('FR-014', 'Accessibility', 'Must', 'System shall pass WCAG 2.1 AA automated axe-core scans on all pages'),
        ('FR-015', 'Security', 'Must', 'System shall return HTTP 401 for missing auth and HTTP 403 for insufficient role'),
    ]
    add_table(doc,
        ['ID', 'Module', 'Priority', 'Requirement'],
        reqs,
        col_widths=[Inches(0.8), Inches(0.9), Inches(0.7), Inches(4.1)])

    heading(doc, '6. Product Roadmap', level=1)
    add_image(doc, 'diag_08_roadmap.png', width=Inches(6.5),
              caption_text='Figure 1 — Five-phase product roadmap')

    heading(doc, '7. Acceptance Criteria', level=1)
    criteria = [
        'Login with valid credentials returns JWT token and redirects to dashboard',
        'Dealer A cannot view Dealer B data even with a valid token',
        'OEM user can access fleet analytics; dealer cannot (403 response)',
        'All 8 predictive models return a risk score for every vehicle in the fleet',
        'Alert table filters correctly by severity and time window',
        'All pages pass axe-core WCAG 2.1 AA scan with zero critical/serious violations',
        'Model retrain completes successfully and logs entry in retrain history',
        'Admin can create a new user who can immediately log in',
    ]
    for c in criteria:
        numbered(doc, c)

    doc.save(DOCS / '03_Product_Requirements.docx')
    print('  ✓ 03_Product_Requirements.docx')


# ── Doc 4 — User Journey ──────────────────────────────────────────────────────
def doc_user_journey():
    doc = new_doc()
    cover_page(doc, 'User Journey Document',
               'Three-persona journey maps with touchpoints, emotions, and pain points')

    heading(doc, '1. Overview')
    body(doc, 'This document maps the end-to-end journeys of the three primary AutoPredict user '
         'personas. Each journey covers the complete flow from initial system access through to '
         'the key value-generating action, including emotional states and system touchpoints.')

    add_image(doc, 'diag_07_user_journey.png', width=Inches(6.5),
              caption_text='Figure 1 — Dealer user journey swimlane (Rajesh Kumar)')

    heading(doc, '2. Persona 1 — Rajesh Kumar (Dealer Service Manager)', level=1)
    heading(doc, 'Journey: From morning login to proactive service booking', level=2)
    steps_rajesh = [
        ('Step 1 — Login', 'Rajesh opens the AutoPredict portal URL and enters credentials.',
         'Curious', 'Login page loads instantly; branded MG Motor layout'),
        ('Step 2 — Dashboard Review', 'Dashboard shows 4 KPI cards. Rajesh sees "3 Critical Alerts" in red.',
         'Focused', 'Fleet Health Score at 87/100; vehicle table sorted by risk'),
        ('Step 3 — Alert Investigation', 'Rajesh clicks to Alerts page; filters to Critical severity.',
         'Concerned', 'Table shows VIN MH01AB1234 — brake pad failure in 12 days'),
        ('Step 4 — Vehicle Detail', 'Rajesh inspects the vehicle record and sees brake pad RUL = 980 km.',
         'Informed', 'Model confidence 91%; last service 8 months ago'),
        ('Step 5 — Service Bay Booking', 'Rajesh opens Service Bay and books Bay 3 for the vehicle.',
         'Confident', 'Booking confirmed; customer notification queued'),
        ('Step 6 — Confirmation', 'Bay occupancy grid reflects the new booking.',
         'Satisfied', 'Proactive appointment booked 12 days before predicted failure'),
    ]
    for step, action, emotion, system_response in steps_rajesh:
        p = doc.add_paragraph()
        r = p.add_run(step + ':  ')
        r.font.bold = True; r.font.color.rgb = BLUE; r.font.size = Pt(10.5)
        p.add_run(action).font.size = Pt(10.5)
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f'Emotion: {emotion}  |  System: ')
        r2.font.italic = True; r2.font.color.rgb = GRAY; r2.font.size = Pt(9.5)
        p2.add_run(system_response).font.size = Pt(9.5)

    heading(doc, '3. Persona 2 — Priya Sharma (OEM Fleet Engineer)', level=1)
    heading(doc, 'Journey: Monthly model health review and retraining', level=2)
    steps_priya = [
        ('Step 1 — OEM Login', 'Priya logs in with OEM credentials; sees expanded navigation including OEM sections.', 'Prepared'),
        ('Step 2 — Fleet Overview', 'Priya views fleet-wide stats grouped by vehicle model. Sees ZS EV has higher alert rate.', 'Analytical'),
        ('Step 3 — Model Health', 'Priya opens Model Health page. Sees hv_battery_soh PSI = 0.28 (threshold: 0.25) in red.', 'Alert'),
        ('Step 4 — EDA Investigation', 'Priya runs EDA on HV battery features. Distribution shift visible in cell_temp histogram.', 'Investigating'),
        ('Step 5 — What-If Analysis', 'Priya tests model with edge-case feature values to understand sensitivity.', 'Exploring'),
        ('Step 6 — Trigger Retrain', 'Priya selects hv_battery_soh and ev_battery_12v, adds training notes, submits retrain.', 'Decisive'),
        ('Step 7 — Verify Results', 'Retrain history shows new AUROC 0.94 (up from 0.87). Priya promotes to champion.', 'Satisfied'),
    ]
    for step, action, emotion in steps_priya:
        p = doc.add_paragraph()
        r = p.add_run(step + ':  ')
        r.font.bold = True; r.font.color.rgb = ORG; r.font.size = Pt(10.5)
        r2 = p.add_run(action)
        r2.font.size = Pt(10.5)
        p2 = doc.add_paragraph()
        r3 = p2.add_run(f'Emotion: {emotion}')
        r3.font.italic = True; r3.font.color.rgb = GRAY; r3.font.size = Pt(9.5)

    heading(doc, '4. Persona 3 — Arjun Mehta (System Admin)', level=1)
    heading(doc, 'Journey: Onboarding a new dealer', level=2)
    steps_arjun = [
        ('Step 1 — Admin Login', 'Arjun logs in with admin credentials.'),
        ('Step 2 — User List', 'Admin panel shows all registered users. Arjun searches for the new dealership.'),
        ('Step 3 — Create User', 'Arjun fills: username, password, role=DEALER, dealer_code=DL099.'),
        ('Step 4 — Confirmation', 'New user appears in list with status Active. Arjun sends credentials to new dealer.'),
        ('Step 5 — Verify Access', 'Arjun logs out and logs in as new dealer to verify correct scoping.'),
    ]
    for step, action in steps_arjun:
        p = doc.add_paragraph()
        r = p.add_run(step + ':  ')
        r.font.bold = True; r.font.color.rgb = GREEN; r.font.size = Pt(10.5)
        p.add_run(action).font.size = Pt(10.5)

    heading(doc, '5. Pain Points & Opportunities', level=1)
    add_table(doc,
        ['Persona', 'Pain Point', 'AutoPredict Solution'],
        [
            ('Rajesh', 'No warning until customer complains', 'ML alerts 30 days before failure'),
            ('Rajesh', 'Bay scheduling is ad hoc', 'Service Bay booking linked to predicted alerts'),
            ('Rajesh', 'Parts often out of stock', 'Inventory demand forecast enables pre-stocking'),
            ('Priya', 'No central model quality view', 'Model Health dashboard with AUROC, PSI per model'),
            ('Priya', 'Retraining requires data science team', 'One-click retrain from OEM portal'),
            ('Arjun', 'User management is manual', 'Admin panel with create/list/deactivate'),
        ])

    doc.save(DOCS / '04_User_Journey.docx')
    print('  ✓ 04_User_Journey.docx')


# ── Doc 5 — Process, Product & Roadmap ───────────────────────────────────────
def doc_roadmap():
    doc = new_doc()
    cover_page(doc, 'Product, Process & Roadmap Document',
               'Core process flows, information architecture, and five-phase delivery plan')

    heading(doc, '1. Core Process Flows')
    body(doc, 'AutoPredict implements four primary operational processes. Each is described below '
         'with inputs, actors, system steps, and outputs.')

    heading(doc, '1.1 Telemetry-to-Alert Process', level=2)
    add_table(doc,
        ['Stage', 'Actor', 'Action', 'Output'],
        [
            ('Signal Emission', 'TBox ECU', 'Emits 200+ signals at 5 Hz over MQTT', 'Raw telemetry stream'),
            ('Ingestion', 'MQTT Broker', 'Receives and queues messages', 'Buffered signal packets'),
            ('Persistence', 'CSV Writer', 'Appends signals to vehicle-partitioned CSV', 'Structured telemetry CSV'),
            ('Feature Engineering', 'FastAPI backend', 'Computes rolling windows, lag features, ratios', 'Feature vector per VIN'),
            ('ML Inference', 'LightGBM models', 'Predicts RUL for 8 components', 'Risk scores + confidence'),
            ('Alert Generation', 'Alert engine', 'Fires alert when RUL < threshold', 'Alert record in alerts CSV'),
            ('UI Display', 'React frontend', 'Fetches and renders alerts table', 'Dealer sees actionable alert'),
        ])

    heading(doc, '1.2 Dealer Predictive Maintenance Process', level=2)
    steps_pm = [
        'Dealer logs in; dashboard loads with fleet risk overview',
        'Dealer identifies high-risk vehicles from KPI cards or vehicle table',
        'Dealer navigates to Alerts to see component-level failure predictions',
        'Dealer opens Service Bay; selects bay; creates appointment for vehicle',
        'Dealer verifies parts availability in Inventory before confirming appointment',
        'On appointment day, technician performs service using predicted component data',
        'Service record updated; ML model receives new ground truth for future training',
    ]
    for s in steps_pm:
        numbered(doc, s)

    heading(doc, '1.3 OEM Model Retraining Process', level=2)
    add_image(doc, 'diag_05_ml_pipeline.png', width=Inches(6.3),
              caption_text='Figure 1 — ML pipeline with champion-challenger model promotion')

    heading(doc, '1.4 Admin User Onboarding Process', level=2)
    for s in ['Admin logs in to admin panel',
              'Admin creates new dealer user with dealer_code',
              'System generates credentials and stores user record',
              'Admin communicates credentials to new dealer contact',
              'Dealer logs in; system applies dealer_code scope to all queries',
              'Admin monitors user list for activity']:
        numbered(doc, s)

    heading(doc, '2. Product Roadmap', level=1)
    add_image(doc, 'diag_08_roadmap.png', width=Inches(6.5),
              caption_text='Figure 2 — Five-phase product delivery roadmap')

    phases_detail = [
        ('Phase 1 — Foundation (Q1 2024) — COMPLETE', '#14A44D', [
            'FastAPI backend with 6 routers and 36 endpoints',
            'JWT authentication with DEALER/OEM/ADMIN roles',
            'CSV-based offline-first data storage',
            'Basic dealer dashboard and alert views',
            'React + TypeScript + Vite frontend scaffold',
        ]),
        ('Phase 2 — ML Core (Q2 2024) — COMPLETE', '#1E5799', [
            '8 LightGBM maintenance prediction models',
            'Feature engineering pipeline (rolling windows, lag features)',
            'Champion-Challenger model registry with PSI monitoring',
            'Alert generation engine with severity thresholds',
            'Automated model evaluation (AUROC, F1, Precision)',
        ]),
        ('Phase 3 — OEM Portal (Q3 2024) — IN PROGRESS', '#E76F51', [
            'OEM Fleet Overview with multi-dimensional grouping',
            'Model Health dashboard with drift indicators',
            'EDA tools (feature distributions, heatmaps)',
            'What-If analysis interface',
            'EV/PHEV health physics engines (DCDC, motor, thermal)',
            'Inventory management with demand forecasting',
        ]),
        ('Phase 4 — Enterprise (Q4 2024) — PLANNED', '#7B2D8B', [
            'PostgreSQL migration (replace CSV with proper database)',
            'Real-time WebSocket push notifications',
            'Multi-worker Uvicorn (horizontal scaling)',
            'Mobile PWA for dealer technicians',
            'Security hardening (bcrypt passwords, rotation key, CORS fix)',
        ]),
        ('Phase 5 — Scale (Q1 2025) — V2', '#0D7377', [
            'Federated ML across multiple OEM brands',
            'Edge inference on TBox hardware',
            'Third-party DMS integration APIs',
            'ISO 21434 cybersecurity certification',
            'Multi-language UI support',
        ]),
    ]
    for phase_name, color_hex, items in phases_detail:
        p = doc.add_paragraph()
        r = p.add_run(phase_name)
        r.font.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(int(color_hex[1:3], 16),
                                     int(color_hex[3:5], 16),
                                     int(color_hex[5:7], 16))
        for item in items:
            bullet(doc, item)

    heading(doc, '3. KPI Framework', level=1)
    add_table(doc,
        ['Metric', 'Baseline', 'Phase 3 Target', 'Phase 4 Target'],
        [
            ('Unplanned breakdowns per 1,000 vehicles', '42/month', '30/month (-29%)', '20/month (-52%)'),
            ('Service bay utilisation', '64%', '75%', '82%'),
            ('Alert-to-booking conversion rate', 'N/A', '45%', '65%'),
            ('ML model AUROC (average)', '0.81', '0.88', '0.92'),
            ('Dealer portal adoption', '0%', '80% of dealers', '95% of dealers'),
            ('Mean time to detect failure', '0 days (reactive)', '30 days', '45 days'),
        ])

    doc.save(DOCS / '05_Product_Process_Roadmap.docx')
    print('  ✓ 05_Product_Process_Roadmap.docx')


# ── Doc 6 — DevOps ────────────────────────────────────────────────────────────
def doc_devops():
    doc = new_doc()
    cover_page(doc, 'DevOps Documentation',
               'Setup, deployment, Docker, CI/CD pipeline, and ML operations')

    heading(doc, '1. System Topology')
    add_image(doc, 'diag_04_deployment.png', width=Inches(6.5),
              caption_text='Figure 1 — Three-container deployment architecture')
    add_table(doc,
        ['Service', 'Port', 'Technology', 'Health Check'],
        [
            ('frontend', '3000', 'React + Vite dev / Nginx prod', 'GET http://localhost:3000/'),
            ('backend', '8001', 'FastAPI + Uvicorn', 'GET http://localhost:8001/health'),
            ('mqtt-broker', '1883', 'Eclipse Mosquitto', 'mosquitto_pub -t test -m ping'),
        ])

    heading(doc, '2. Development Setup', level=1)
    heading(doc, '2.1 Prerequisites', level=2)
    bullet(doc, 'Python 3.11+ with pip')
    bullet(doc, 'Node.js 20 LTS with npm')
    bullet(doc, 'Git 2.40+')
    bullet(doc, 'Eclipse Mosquitto (optional for local MQTT testing)')

    heading(doc, '2.2 Backend Setup', level=2)
    for cmd in [
        'cd api && pip install -r requirements.txt',
        'cp .env.example .env  # then edit SECRET_KEY',
        'uvicorn api.main:app --reload --port 8001',
    ]:
        p = doc.add_paragraph(cmd, style='No Spacing')
        for run in p.runs:
            run.font.name = 'Courier New'; run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
            p.paragraph_format.left_indent = Cm(1)

    heading(doc, '2.3 Frontend Setup', level=2)
    for cmd in [
        'cd dealer_portal && npm install',
        'npm run dev  # starts Vite at http://localhost:3000',
    ]:
        p = doc.add_paragraph(cmd, style='No Spacing')
        for run in p.runs:
            run.font.name = 'Courier New'; run.font.size = Pt(9.5)
            p.paragraph_format.left_indent = Cm(1)

    heading(doc, '3. Environment Variables', level=1)
    add_table(doc,
        ['Variable', 'Required', 'Default', 'Description'],
        [
            ('SECRET_KEY', 'CRITICAL', 'change-me-in-production', 'JWT signing key — MUST be replaced in production (32+ random bytes)'),
            ('CORS_ORIGINS', 'Required in prod', '["*"]', 'Comma-separated list of allowed frontend origins'),
            ('MQTT_BROKER', 'Optional', 'localhost', 'MQTT broker hostname for TBox telemetry'),
            ('MQTT_PORT', 'Optional', '1883', 'MQTT broker port'),
            ('MQTT_TOPIC', 'Optional', 'mg/vehicles/#', 'MQTT subscription topic'),
            ('DATA_DIR', 'Optional', './data', 'Directory containing CSV data files'),
            ('MODEL_DIR', 'Optional', './models', 'Directory containing trained .pkl model artifacts'),
        ],
        col_widths=[Inches(1.3), Inches(0.9), Inches(1.5), Inches(2.8)])

    heading(doc, '4. Docker Deployment', level=1)
    for line in [
        'docker-compose build    # build all three images',
        'docker-compose up -d    # start in detached mode',
        'docker-compose logs -f  # tail all service logs',
        'docker-compose down     # stop all services',
    ]:
        p = doc.add_paragraph(line, style='No Spacing')
        for run in p.runs:
            run.font.name = 'Courier New'; run.font.size = Pt(9.5)
            p.paragraph_format.left_indent = Cm(1)

    heading(doc, '5. CI/CD Pipeline', level=1)
    ci_stages = [
        ('Checkout', 'actions/checkout@v4'),
        ('Python Setup', 'actions/setup-python@v5 (3.11)'),
        ('Backend Tests', 'pytest api/tests/ -v --cov'),
        ('Node Setup', 'actions/setup-node@v4 (Node 20)'),
        ('Frontend Build', 'npm ci && npm run build'),
        ('Playwright Install', 'npx playwright install --with-deps chromium'),
        ('E2E Tests', 'npx playwright test (workers=1)'),
        ('Report Upload', 'Upload playwright-report/ artifact'),
    ]
    for stage, command in ci_stages:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{stage}: ')
        r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
        p.add_run(command).font.size = Pt(10)

    heading(doc, '6. ML Operations', level=1)
    add_image(doc, 'diag_05_ml_pipeline.png', width=Inches(6.3),
              caption_text='Figure 2 — ML pipeline and champion-challenger model management')
    body(doc, 'Model artifacts are stored in the models/ directory as .pkl files. Each model '
         'file is versioned with a timestamp suffix. The ModelRegistry class manages champion '
         'promotion and PSI drift calculation.')
    add_table(doc,
        ['Model File', 'Retrain Command', 'Typical AUROC'],
        [
            ('brake_pad.pkl', 'py train_all.py --model brake_pad', '0.91'),
            ('engine_oil.pkl', 'py train_all.py --model engine_oil', '0.88'),
            ('tyre_wear.pkl', 'py train_all.py --model tyre_wear', '0.86'),
            ('battery_12v.pkl', 'py train_all.py --model battery_12v', '0.89'),
            ('hv_battery_soh.pkl', 'py train_all.py --model hv_battery_soh', '0.93'),
            ('transmission.pkl', 'py train_all.py --model transmission', '0.85'),
            ('cooling_system.pkl', 'py train_all.py --model cooling_system', '0.87'),
            ('driver_score.pkl', 'py train_all.py --model driver_score', '0.84'),
        ])

    heading(doc, '7. Production Security Checklist', level=1)
    checklist = [
        'Replace SECRET_KEY with 32-byte random value: python -c "import secrets; print(secrets.token_hex(32))"',
        'Implement bcrypt password hashing (passlib already installed; call verify_password() on login)',
        'Set CORS_ORIGINS to specific frontend domain — remove wildcard ["*"]',
        'Add JWT sub validation against user store in get_current_user()',
        'Enable HTTPS (TLS certificate via Let\'s Encrypt or corporate CA)',
        'Configure Nginx rate limiting in addition to slowapi',
        'Rotate SECRET_KEY every 90 days with key overlap period for in-flight tokens',
        'Enable Mosquitto TLS and MQTT authentication',
        'Set up log rotation and monitoring alerts for auth failures',
        'Run OWASP ZAP scan against staging before production deployment',
    ]
    for item in checklist:
        bullet(doc, item)

    doc.save(DOCS / '06_DevOps_Documentation.docx')
    print('  ✓ 06_DevOps_Documentation.docx')


# ── Doc 7 — Frontend ──────────────────────────────────────────────────────────
def doc_frontend():
    doc = new_doc()
    cover_page(doc, 'Frontend Documentation',
               'React architecture, component structure, routing, state management, and accessibility')

    heading(doc, '1. Technology Stack')
    add_table(doc,
        ['Technology', 'Version', 'Role', 'Key Decision'],
        [
            ('React', '18', 'UI framework', 'Concurrent features; large ecosystem'),
            ('TypeScript', '5', 'Type safety', 'Catches prop/API contract errors at compile time'),
            ('Vite', '5', 'Build tool', 'Sub-second HMR; native ESM; faster than Webpack'),
            ('Tailwind CSS', '3', 'Styling', 'Utility-first; zero runtime; PurgeCSS in prod'),
            ('TanStack Query', 'v5', 'Server state', 'Auto-caching; background refetch; stale-while-revalidate'),
            ('Axios', '1.x', 'HTTP client', 'Interceptors for token injection and 401 handling'),
            ('React Router', 'v6', 'Routing', 'PrivateRoute + OemRoute guards; nested routes'),
            ('Recharts', '2', 'Charts', 'SVG-based; props passed to path elements for aria-label'),
        ])

    heading(doc, '2. Application Routing', level=1)
    routes = [
        ('/login', 'LoginPage', 'Public', 'Redirects to / if already authenticated'),
        ('/', 'DashboardPage', 'DEALER', 'Fleet health overview with vehicle table'),
        ('/alerts', 'AlertsPage', 'DEALER', 'Predictive alerts with severity and time filters'),
        ('/service', 'ServiceBayPage', 'DEALER', 'Bay grid and appointment booking'),
        ('/inventory', 'InventoryPage', 'DEALER', '7-tab inventory management'),
        ('/oem/fleet', 'OemFleetOverview', 'OEM', 'Fleet-wide analytics with group-by controls'),
        ('/oem/models', 'OemModelHealth', 'OEM', 'AUROC/PSI model quality dashboard'),
        ('/oem/eda', 'OemEda', 'OEM', 'Exploratory data analysis tools'),
        ('/oem/whatif', 'OemWhatIf', 'OEM', 'Scenario analysis for model predictions'),
        ('/oem/retrain', 'OemRetrain', 'OEM', 'Model retrain trigger and history'),
        ('/admin', 'AdminPanel', 'ADMIN', 'User management (CRUD)'),
    ]
    add_table(doc,
        ['Route', 'Component', 'Auth Required', 'Description'],
        routes,
        col_widths=[Inches(1.3), Inches(1.5), Inches(0.8), Inches(2.9)])

    heading(doc, '3. Authentication Flow', level=1)
    add_image(doc, 'diag_03_auth_rbac.png', width=Inches(6.3),
              caption_text='Figure 1 — Auth flow and RBAC hierarchy')
    body(doc, 'Auth state is stored in localStorage with four keys:')
    add_table(doc,
        ['Key', 'Value', 'Purpose'],
        [
            ('ap_token', 'JWT string', 'Sent as Authorization: Bearer header'),
            ('ap_role', 'dealer / oem / admin', 'Controls route access and sidebar items'),
            ('ap_dealer_code', 'DL001, DL002 …', 'Scopes all API queries to this dealer'),
            ('ap_user', 'username string', 'Displayed in sidebar'),
        ])

    heading(doc, '4. State Management', level=1)
    body(doc, 'State is divided into two clear categories:')
    bullet(doc, 'Server state (API data): TanStack Query — caching, background refetch, loading/error states')
    bullet(doc, 'Client state (UI state): React useState and Context — sidebar toggle, filter selections, modal open/close')
    body(doc, 'No Redux or global state manager is used. The separation of concerns is enforced at the '
         'component level: components use useQuery() for data and useState() for UI.')

    heading(doc, '5. Key Pages', level=1)
    pages_detail = [
        ('DashboardPage', 'src/pages/Dashboard.tsx', ['4 KPI stat cards (Total Vehicles, Active Alerts, Due Service, Health Score)', 'Searchable, sortable vehicle table (VIN, Model, Year, Mileage, Risk, Status)', 'Upcoming service panel (next 7 days)', 'Real-time refresh via TanStack Query (5 min stale time)']),
        ('AlertsPage', 'src/pages/Alerts.tsx', ['Severity filter buttons (All/Critical/High/Medium/Low) with active state highlighting', 'Time window selector (24h/7d/30d/90d)', 'Manual refresh with loading spinner', 'Empty state illustration when no alerts match']),
        ('OemFleetOverview', 'src/pages/OemFleetOverview.tsx', ['Group-by selector: Model / Region / Year / Fuel Type', 'Bar chart of fleet health by group (Recharts)', 'Pie chart of alert severity distribution', 'Summary KPI cards with fleet-wide totals']),
        ('OemRetrain', 'src/pages/OemRetrain.tsx', ['Checkbox list for 8 LightGBM models + 4 EV engines', 'Select All / Clear All buttons', 'Training notes textarea', 'Retrain history table with AUROC before/after']),
    ]
    for page_name, filepath, features in pages_detail:
        heading(doc, f'{page_name} ({filepath})', level=2)
        for f in features:
            bullet(doc, f)

    heading(doc, '6. Accessibility (WCAG 2.1 AA)', level=1)
    add_table(doc,
        ['Issue', 'Component', 'Fix Applied'],
        [
            ('svg-img-alt', 'OemFleetOverview Recharts PieChart', 'Added aria-label to each <Cell>: "{name}: {value} vehicles"'),
            ('scrollable-region-focusable', 'OemModelHealth heatmap div', 'Added tabIndex={0} to overflow-auto container'),
        ])
    body(doc, 'All 8 pages pass axe-core automated scans with zero critical or serious violations. '
         'The Playwright test suite (10-accessibility.spec.ts) runs these scans on every CI run.')

    doc.save(DOCS / '07_Frontend_Documentation.docx')
    print('  ✓ 07_Frontend_Documentation.docx')


# ── Doc 8 — Backend API ───────────────────────────────────────────────────────
def doc_backend():
    doc = new_doc()
    cover_page(doc, 'Backend API Documentation',
               'FastAPI endpoints, authentication, request/response schemas, and error codes')

    heading(doc, '1. API Overview')
    add_image(doc, 'diag_06_api_routes.png', width=Inches(6.5),
              caption_text='Figure 1 — API router map: 36 endpoints across 6 routers')
    body(doc, 'The AutoPredict backend is a FastAPI application exposing 36 REST endpoints '
         'across 6 routers. All non-public endpoints require JWT Bearer authentication. '
         'The base URL is http://localhost:8001 in development.')

    heading(doc, '2. Authentication', level=1)
    add_image(doc, 'diag_03_auth_rbac.png', width=Inches(6.3),
              caption_text='Figure 2 — Authentication and RBAC flow')

    heading(doc, '2.1 Obtain Token', level=2)
    body(doc, 'POST /api/auth/token')
    add_table(doc,
        ['Field', 'Request Body', 'Response'],
        [
            ('username', 'string (required)', '—'),
            ('password', 'string (required)', '—'),
            ('access_token', '—', 'JWT string'),
            ('role', '—', 'dealer / oem / admin'),
            ('dealer_code', '—', 'DL001 / null'),
            ('token_type', '—', '"bearer"'),
        ])

    heading(doc, '3. Routers & Endpoints', level=1)

    routers = [
        ('/api/auth', 'Public', [
            ('POST', '/token', 'Authenticate user; returns JWT token'),
            ('GET', '/me', 'Returns current user info from JWT claims'),
        ]),
        ('/api/fleet', 'DEALER', [
            ('GET', '/dashboard', 'Fleet KPI summary (total vehicles, alerts, health score)'),
            ('GET', '/vehicles', 'Paginated vehicle list with risk scores (scoped to dealer_code)'),
            ('GET', '/vehicles/{vin}', 'Single vehicle detail with all ML predictions'),
            ('GET', '/alerts', 'Alert records with severity filter and time window'),
            ('GET', '/upcoming-service', 'Vehicles due for service in next N days'),
            ('GET', '/health-trend', 'Fleet health score time series'),
        ]),
        ('/api/dealer', 'DEALER', [
            ('GET', '/service-bay', 'Bay grid status and appointment list'),
            ('POST', '/service-bay/book', 'Create bay appointment (VIN, bay, date, type)'),
            ('GET', '/demand-forecast', 'ML-driven parts demand forecast'),
            ('GET', '/inventory', 'Stock levels with reorder flags'),
        ]),
        ('/api/oem', 'OEM', [
            ('GET', '/fleet-overview', 'Fleet-wide stats grouped by model/region/year/fuel'),
            ('GET', '/model-health', 'AUROC, PSI, precision, recall per model per region'),
            ('GET', '/eda', 'Feature distributions and correlation for EDA'),
            ('POST', '/whatif', 'What-if prediction: POST feature values, GET prediction'),
            ('POST', '/retrain', 'Trigger model retraining for selected models'),
            ('GET', '/retrain/history', 'Retrain log with AUROC before/after'),
        ]),
        ('/api/admin', 'ADMIN', [
            ('GET', '/users', 'List all registered users'),
            ('POST', '/users', 'Create new user (username, password, role, dealer_code)'),
            ('DELETE', '/users/{username}', 'Deactivate user account'),
            ('PUT', '/users/{username}', 'Update user role or dealer_code'),
            ('GET', '/audit-log', 'Admin action audit trail'),
        ]),
    ]
    for prefix, role, endpoints in routers:
        heading(doc, f'{prefix}  (Role: {role})', level=2)
        add_table(doc,
            ['Method', 'Path', 'Description'],
            [(m, p, d) for m, p, d in endpoints],
            col_widths=[Inches(0.7), Inches(2.0), Inches(3.8)])
        spacer(doc)

    heading(doc, '4. Error Codes', level=1)
    add_table(doc,
        ['HTTP Code', 'Condition', 'Response Body'],
        [
            ('200', 'Success', '{"data": ...}'),
            ('400', 'Invalid request body (Pydantic validation)', '{"detail": "field required"}'),
            ('401', 'Missing or malformed Authorization header', '{"detail": "Not authenticated"}'),
            ('403', 'Valid token; insufficient role', '{"detail": "Forbidden"}'),
            ('404', 'Resource not found (VIN, user, etc.)', '{"detail": "Not found"}'),
            ('422', 'Unprocessable entity (FastAPI Pydantic error)', '{"detail": [{...}]}'),
            ('429', 'Rate limit exceeded (auth endpoints)', '{"detail": "Rate limit exceeded"}'),
            ('500', 'Internal server error (CSV parse failure, model error)', '{"detail": "Internal error"}'),
        ])

    heading(doc, '5. Data Models', level=1)
    add_image(doc, 'diag_02_data_flow.png', width=Inches(6.3),
              caption_text='Figure 3 — Data flow from TBox signal to API response')

    doc.save(DOCS / '08_Backend_API_Documentation.docx')
    print('  ✓ 08_Backend_API_Documentation.docx')


# ── Doc 9 — Test Cases ────────────────────────────────────────────────────────
def doc_test_cases():
    doc = new_doc()
    cover_page(doc, 'Test Cases Document',
               '120 test cases across 9 modules — functional, automated, and manual')

    heading(doc, '1. Test Coverage Overview')
    add_table(doc,
        ['Module', 'Test IDs', 'Count', 'Automated', 'Manual'],
        [
            ('TC-01 Authentication', 'TC-001 – TC-012', '12', '10', '2'),
            ('TC-02 Fleet Dashboard', 'TC-013 – TC-022', '10', '7', '3'),
            ('TC-03 Alerts', 'TC-023 – TC-032', '10', '8', '2'),
            ('TC-04 Service Bay', 'TC-033 – TC-042', '10', '4', '6'),
            ('TC-05 Inventory', 'TC-043 – TC-056', '14', '6', '8'),
            ('TC-06 OEM Fleet', 'TC-057 – TC-066', '10', '8', '2'),
            ('TC-07 OEM Retrain', 'TC-067 – TC-076', '10', '10', '0'),
            ('TC-08 Admin', 'TC-077 – TC-086', '10', '5', '5'),
            ('TC-09 Accessibility', 'TC-087 – TC-100', '14', '14', '0'),
            ('TOTAL', '', '120', '89', '31'),
        ])

    heading(doc, '2. Authentication Test Cases (TC-01)', level=1)
    auth_cases = [
        ('TC-001', 'Valid login — dealer role', 'P1', 'AUTO',
         'POST /api/auth/token with valid dealer credentials',
         'HTTP 200; JWT returned; role=dealer; dealer_code present'),
        ('TC-002', 'Valid login — OEM role', 'P1', 'AUTO',
         'POST /api/auth/token with OEM credentials',
         'HTTP 200; role=oem; dealer_code=null'),
        ('TC-003', 'Invalid password', 'P1', 'AUTO',
         'POST /api/auth/token with wrong password',
         'HTTP 401; {"detail": "Invalid credentials"}'),
        ('TC-004', 'Missing authorization header', 'P1', 'AUTO',
         'GET /api/fleet/dashboard without Authorization header',
         'HTTP 401; {"detail": "Not authenticated"}'),
        ('TC-005', 'Expired JWT', 'P1', 'AUTO',
         'Send JWT with exp in the past',
         'HTTP 401; {"detail": "Token expired"}'),
        ('TC-006', 'Tampered JWT signature', 'P1', 'AUTO',
         'Modify JWT payload without re-signing',
         'HTTP 401; {"detail": "Invalid token"}'),
        ('TC-007', 'Dealer accesses OEM endpoint', 'P1', 'AUTO',
         'GET /api/oem/fleet-overview with dealer token',
         'HTTP 403; {"detail": "Forbidden"}'),
        ('TC-008', 'Redirect to login when unauthenticated', 'P1', 'AUTO',
         'Navigate to / without auth in localStorage',
         'Browser redirects to /login'),
        ('TC-009', 'Login form validation — empty fields', 'P2', 'AUTO',
         'Submit login form with empty username and password',
         'Form shows validation error; no API call made'),
        ('TC-010', 'Sign out clears auth state', 'P1', 'AUTO',
         'Click Sign Out button',
         'localStorage cleared; redirected to /login'),
        ('TC-011', 'Admin creates user via API', 'P1', 'AUTO',
         'POST /api/admin/users with admin token and new user payload',
         'HTTP 201; new user can immediately authenticate'),
        ('TC-012', 'Rate limiting on auth endpoint', 'P2', 'MANUAL',
         'Send 6+ rapid requests to /api/auth/token',
         '6th request returns HTTP 429'),
    ]
    add_table(doc,
        ['ID', 'Test Name', 'Priority', 'Type', 'Steps', 'Expected'],
        auth_cases,
        col_widths=[Inches(0.65), Inches(1.5), Inches(0.55), Inches(0.55), Inches(2.0), Inches(1.6)])

    heading(doc, '3. Fleet Dashboard Test Cases (TC-02)', level=1)
    dashboard_cases = [
        ('TC-013', 'Dashboard loads for dealer', 'P1', 'AUTO', 'Log in as dealer; navigate to /', 'Heading "Fleet Dashboard" visible; 4 KPI cards rendered'),
        ('TC-014', 'KPI cards display values', 'P1', 'AUTO', 'Wait for data to load', 'Total Vehicles, Active Alerts, Due Service, Health Score all show non-empty values'),
        ('TC-015', 'Vehicle table renders', 'P1', 'AUTO', 'Wait for table to appear', 'Table visible with columns: VIN, Model, Year, Risk Score, Status'),
        ('TC-016', 'Search filters table', 'P2', 'AUTO', 'Type VIN into search box', 'Table rows filter to matching VINs only'),
        ('TC-017', 'Sort by Risk Score', 'P2', 'AUTO', 'Click Risk Score column header', 'Rows reorder by descending risk score'),
        ('TC-018', 'OEM cannot access via /api/fleet', 'P1', 'AUTO', 'Send GET /api/fleet/dashboard with OEM token', 'HTTP 200 (OEM has read access to all fleet data)'),
        ('TC-019', 'Dashboard scoped to dealer_code', 'P1', 'MANUAL', 'Log in as DL001; verify VINs shown belong to DL001 only', 'No DL002 vehicles appear in table'),
        ('TC-020', 'Upcoming service panel', 'P2', 'AUTO', 'Wait for data; check right panel', 'Upcoming service list shows appointments in next 7 days'),
        ('TC-021', 'Sign out from dashboard', 'P1', 'AUTO', 'Click user menu; click Sign Out', 'Auth cleared; redirected to /login'),
        ('TC-022', 'Dashboard not accessible when unauthenticated', 'P1', 'AUTO', 'Clear localStorage; navigate to /', 'Redirected to /login'),
    ]
    add_table(doc,
        ['ID', 'Test Name', 'Priority', 'Type', 'Steps', 'Expected'],
        dashboard_cases,
        col_widths=[Inches(0.65), Inches(1.5), Inches(0.55), Inches(0.55), Inches(2.0), Inches(1.6)])

    heading(doc, '4. Accessibility Test Cases (TC-09)', level=1)
    a11y_cases = [
        ('TC-087', 'Login page WCAG 2.1 AA', 'P1', 'AUTO', 'Run axe-core on /login', 'Zero critical or serious violations'),
        ('TC-088', 'Dashboard WCAG 2.1 AA', 'P1', 'AUTO', 'Run axe-core on / after data loads', 'Zero critical or serious violations'),
        ('TC-089', 'Alerts page WCAG 2.1 AA', 'P1', 'AUTO', 'Run axe-core on /alerts', 'Zero critical or serious violations'),
        ('TC-090', 'OEM Fleet WCAG 2.1 AA', 'P1', 'AUTO', 'Run axe-core on /oem/fleet', 'Zero critical or serious violations; pie chart cells have aria-label'),
        ('TC-091', 'OEM Model Health WCAG 2.1 AA', 'P1', 'AUTO', 'Run axe-core on /oem/models', 'Zero violations; heatmap div has tabIndex=0'),
        ('TC-092', 'OEM Retrain WCAG 2.1 AA', 'P1', 'AUTO', 'Run axe-core on /oem/retrain', 'Zero critical or serious violations'),
        ('TC-093', 'Keyboard navigation — Login', 'P1', 'AUTO', 'Tab through login form; press Enter to submit', 'All interactive elements reachable by Tab; form submits'),
        ('TC-094', 'Keyboard navigation — Dashboard', 'P2', 'AUTO', 'Tab through dashboard', 'Sidebar links, table, sort controls keyboard accessible'),
        ('TC-095', 'Colour contrast — headings', 'P1', 'AUTO', 'axe-core color-contrast rule', 'All text passes WCAG AA contrast ratio (4.5:1 for normal text)'),
        ('TC-096', 'Screen reader — table headers', 'P2', 'MANUAL', 'Use NVDA; navigate vehicle table', 'Column headers announced; cell values associated with headers'),
        ('TC-097', 'Focus visible on all controls', 'P1', 'AUTO', 'Tab through entire page', 'Focus ring visible on all interactive elements'),
        ('TC-098', 'Error messages announced', 'P2', 'MANUAL', 'Submit invalid login; check screen reader', 'Error message "Invalid credentials" announced via live region'),
        ('TC-099', 'Images have alt text', 'P1', 'AUTO', 'axe-core image-alt rule', 'All <img> elements have non-empty alt attribute'),
        ('TC-100', 'No keyboard traps', 'P1', 'MANUAL', 'Open modal; Tab through; press Escape', 'Focus returns to trigger element; Escape closes modal'),
    ]
    add_table(doc,
        ['ID', 'Test Name', 'Priority', 'Type', 'Steps', 'Expected'],
        a11y_cases,
        col_widths=[Inches(0.65), Inches(1.5), Inches(0.55), Inches(0.55), Inches(2.0), Inches(1.6)])

    body(doc, 'Full 120 test cases are available in the HTML artifact (Test Cases Document). '
         'This Word document includes the key test groups for quick reference.')

    doc.save(DOCS / '09_Test_Cases.docx')
    print('  ✓ 09_Test_Cases.docx')


# ── Doc 10 — Playwright ───────────────────────────────────────────────────────
def doc_playwright():
    doc = new_doc()
    cover_page(doc, 'Playwright Test Documentation',
               '174 tests · 12 suites · 100% passing — infrastructure, patterns, and full test inventory')

    heading(doc, '1. Framework Overview')
    body(doc, 'AutoPredict uses Playwright 1.61.1 as the primary end-to-end test framework. '
         'The suite covers 174 tests across 12 spec files, exercising the full stack from '
         'React UI through FastAPI backend to CSV data processing.')

    add_table(doc,
        ['Metric', 'Value'],
        [
            ('Total Tests', '174 (100% passing)'),
            ('Browser', 'Chromium (Playwright bundled)'),
            ('Workers', '1 (mandatory — CSV backend is single-threaded)'),
            ('Global Timeout', '120 seconds (cold CSV reads take 55–65s)'),
            ('Expect Timeout', '12 seconds'),
            ('Navigation Timeout', '25 seconds'),
            ('Auth Strategy', 'Pre-serialised storageState per role (global-setup.ts)'),
            ('Page Objects', '8 POMs (one per major page)'),
            ('Auth Roles', '4 (dealer, dealer2, oem, admin)'),
        ])

    heading(doc, '2. Critical Design Decision: workers=1', level=1)
    body(doc, 'The FastAPI backend processes CSV files synchronously using pandas. '
         'Multiple concurrent Playwright workers would queue backend requests, causing '
         '25–45 second timeouts and flaky tests. The workers:1 setting is mandatory — '
         'not an oversight. All 174 tests run serially in approximately 6 minutes.')

    heading(doc, '3. Data API Regex Pattern', level=1)
    body(doc, 'The most important utility in the test suite is the DATA_API regex:')
    p = doc.add_paragraph('const DATA_API = /https?:\\/\\/[^/]+\\/api\\//;', style='No Spacing')
    for run in p.runs:
        run.font.name = 'Courier New'; run.font.size = Pt(9.5)
        p.paragraph_format.left_indent = Cm(1)
    body(doc, 'This is used as: await page.route(DATA_API, r => r.abort())')
    body(doc, 'WHY regex, not glob: The glob pattern **/api/** also matches the Vite source '
         'file at /src/api/client.ts. When this file is aborted, React fails to mount — '
         'route guards never fire and all navigation tests fail. The regex anchors /api/ '
         'as the first path segment after the host, matching only real backend requests.')

    heading(doc, '4. Authentication Infrastructure', level=1)
    body(doc, 'global-setup.ts runs once before all tests and generates 4 storageState files '
         '(e2e/.auth/*.json). Each file contains the localStorage keys needed for that role. '
         'Tests import the auth fixture (e2e/fixtures/auth.ts) which provides typed page '
         'fixtures: dealerPage, dealer2Page, oemPage, adminPage.')

    heading(doc, '5. Test Suite Inventory', level=1)
    suites = [
        ('01-auth.spec.ts', '13', 'Login form, error handling, redirect guards, sign-out, branding. 3 tests marked test.slow() for 180s login flow timeout.'),
        ('02-dealer-dashboard.spec.ts', '7', 'Dashboard heading, KPI cards, vehicle table, search, sort, upcoming service, sign-out. All abort DATA_API for non-data tests.'),
        ('03-dealer-alerts.spec.ts', '8', 'All 5 filter buttons, hours select, refresh button, data load (waitForData), column headers.'),
        ('06-oem-fleet.spec.ts', '8', 'Group-by all 4 options (model/region/year/fuel), KPI cards, chart sections. No route abort — needs live data.'),
        ('08-oem-retrain.spec.ts', '13', 'Select all, clear all, submit state, notes field, history table, individual model checkboxes, EV engine section.'),
        ('09-role-guards.spec.ts', '17', 'Every protected route tested for each wrong role. All use DATA_API regex abort.'),
        ('10-accessibility.spec.ts', '8', 'axe-core WCAG 2.1 AA scan on 8 pages. OEM tests use 55s heading timeout for cold CSV reads.'),
        ('11-api-contracts.spec.ts', '15', 'Direct HTTP via APIRequestContext — no browser. JWT shape, fleet-overview schema, demand-forecast scope, model-health format.'),
        ('12-security.spec.ts', '59', '8 sections A–H: auth bypass, JWT forgery, RBAC vertical, horizontal isolation, input validation, info disclosure, transport, hardening.'),
        ('global-setup.ts', '4', 'Generates dealer.json, dealer2.json, oem.json, admin.json storageState files.'),
        ('global-teardown.ts', '1', 'Removes all .auth/*.json files after suite completes.'),
    ]
    add_table(doc,
        ['File', 'Tests', 'Coverage'],
        suites,
        col_widths=[Inches(2.2), Inches(0.6), Inches(3.7)])

    heading(doc, '6. Page Object Models', level=1)
    poms = [
        ('LoginPage', 'e2e/pages/LoginPage.ts', ['usernameInput, passwordInput, submitBtn', 'errorMsg, loadingSpinner', 'fill(user, pass), submit(), waitForError()', 'expectRedirectToDashboard()']),
        ('DashboardPage', 'e2e/pages/DashboardPage.ts', ['heading, kpiCards, vehicleTable', 'searchInput, signOutBtn', 'waitForTable()', 'getKpiValue(name)']),
        ('AlertsPage', 'e2e/pages/AlertsPage.ts', ['heading, hoursSelect, refreshBtn', 'table, emptyState', 'filterBtn(severity)', 'setSeverity(), setHours(), waitForData()']),
        ('OemFleetPage', 'e2e/pages/OemFleetPage.ts', ['heading, groupBySelect, charts', 'waitForData() with 80s timeout', 'getGroupByOption(name)']),
        ('OemRetrainPage', 'e2e/pages/OemRetrainPage.ts', ['selectAllBtn, clearBtn, submitBtn', 'notesField, historyTable', 'selectModel(name), submitRetrain()', 'waitForHistory()']),
    ]
    for pom_name, filepath, members in poms:
        heading(doc, f'{pom_name}  ({filepath})', level=2)
        for m in members:
            bullet(doc, m)

    heading(doc, '7. Missing Test Files (Coverage Gaps)', level=1)
    body(doc, 'Three spec files are not yet written. Estimated 30 additional tests required:')
    add_table(doc,
        ['File', 'Coverage Missing', 'Estimated Tests'],
        [
            ('04-service-bay.spec.ts', 'Bay grid, booking modal (open/fill/submit/cancel), appointment list', '10'),
            ('05-inventory.spec.ts', 'All 7 tabs, demand forecast horizon switching, stock levels, EV parts', '12'),
            ('07-oem-models.spec.ts', 'Model cards (AUROC, drift status), EDA sub-tabs, What-If form', '8'),
        ])

    doc.save(DOCS / '10_Playwright_Test_Documentation.docx')
    print('  ✓ 10_Playwright_Test_Documentation.docx')


# ── Doc 11 — Security ─────────────────────────────────────────────────────────
def doc_security():
    doc = new_doc()
    cover_page(doc, 'Security Test Report',
               '59 tests across 8 domains · 5 confirmed vulnerabilities · Remediation roadmap')

    heading(doc, '1. Executive Summary')
    body(doc, 'A comprehensive security test suite was executed against the AutoPredict platform '
         'covering authentication, authorisation, JWT integrity, input validation, information '
         'disclosure, and transport security. 59 tests were run across 8 security domains.')

    body(doc, '5 vulnerabilities were confirmed. All require remediation before production deployment.')
    add_image(doc, 'diag_09_security.png', width=Inches(6.5),
              caption_text='Figure 1 — Security test coverage across 8 domains')

    heading(doc, '2. Confirmed Vulnerabilities', level=1)
    add_table(doc,
        ['ID', 'Severity', 'Title', 'Status'],
        [
            ('VLN-001', 'CRITICAL', 'Default SECRET_KEY in production', 'Open'),
            ('VLN-002', 'CRITICAL', 'Plaintext password storage', 'Open'),
            ('VLN-003', 'HIGH', 'CORS wildcard allows any origin', 'Open'),
            ('VLN-004', 'HIGH', 'JWT sub not validated against user store', 'Open'),
            ('VLN-005', 'HIGH', 'Forged dealer_code grants tenant privilege escalation', 'Open'),
        ])

    vulns_detail = [
        ('VLN-001', 'CRITICAL', 'Default SECRET_KEY',
         'The JWT signing key defaults to "change-me-in-production". Any attacker who reads the '
         'source code (which is likely for an internal platform) can forge valid JWTs for any role.',
         'Generate a 32-byte random key: python -c "import secrets; print(secrets.token_hex(32))". '
         'Add startup guard that raises if SECRET_KEY == "change-me-in-production".'),
        ('VLN-002', 'CRITICAL', 'Plaintext passwords',
         'User passwords are stored in plaintext. verify_password() exists (using passlib/bcrypt) '
         'but is never called. A database breach exposes all credentials immediately.',
         'Call pwd_context.verify(plain, hashed) on login. Hash passwords with '
         'pwd_context.hash() on user creation. Passlib is already installed.'),
        ('VLN-003', 'HIGH', 'CORS wildcard',
         'CORS is configured with allow_origins=["*"]. Any malicious website can make '
         'authenticated requests to the API using a victim\'s credentials.',
         'Set CORS_ORIGINS env var to the specific frontend domain. '
         'Remove the wildcard ["*"] from FastAPI CORSMiddleware configuration.'),
        ('VLN-004', 'HIGH', 'JWT sub not validated',
         'The JWT sub (username) claim is trusted without verifying the user still exists '
         'in the user store. A deleted user\'s token remains valid until expiry.',
         'In get_current_user(), look up users[payload["sub"]] and raise 401 if not found. '
         'This also enables token invalidation on password change.'),
        ('VLN-005', 'HIGH', 'Tenant privilege escalation via forged token',
         'A dealer who forges a JWT with dealer_code=DL002 (using the default SECRET_KEY, VLN-001) '
         'can access another dealer\'s complete vehicle and alert data.',
         'This is a compound vulnerability — fixing VLN-001 (SECRET_KEY) and VLN-002 '
         '(user validation) eliminates the attack surface. Also enforce server-side '
         'dealer_code lookup from user store, not from JWT claims.'),
    ]
    for vid, sev, title, impact, remediation in vulns_detail:
        heading(doc, f'{vid} — {title} [{sev}]', level=2)
        p = doc.add_paragraph()
        r = p.add_run('Impact: ')
        r.font.bold = True; r.font.color.rgb = RED
        p.add_run(impact).font.size = Pt(10.5)
        p = doc.add_paragraph()
        r = p.add_run('Remediation: ')
        r.font.bold = True; r.font.color.rgb = GREEN
        p.add_run(remediation).font.size = Pt(10.5)

    heading(doc, '3. Security Test Sections', level=1)
    sections = [
        ('A', 'Authentication Bypass', 9, [
            'Missing Authorization header → HTTP 401 (not 403)',
            'Empty Authorization header → HTTP 401',
            'Malformed bearer token → HTTP 401',
            'Completely invalid JWT string → HTTP 401',
            'Expired JWT → HTTP 401',
            'JWT with future iat → HTTP 401',
            'JWT with none algorithm → HTTP 401',
            '[VULN] Default SECRET_KEY allows forged token → HTTP 200 (VLN-001)',
            'Correct-format wrong-secret JWT → HTTP 401',
        ]),
        ('B', 'JWT Token Forgery', 4, [
            '[VULN] Forge admin token with default secret → HTTP 200 (VLN-001)',
            'Tampered payload (invalid signature) → HTTP 401',
            'Forged token for non-existent user → HTTP 200 (VLN-004)',
            'Token with modified expiry → HTTP 401',
        ]),
        ('C', 'RBAC — Vertical Isolation', 10, [
            'Dealer cannot access /api/oem/* → HTTP 403',
            'Dealer cannot access /api/admin/* → HTTP 403',
            'OEM cannot access /api/admin/* → HTTP 403',
            'OEM can access /api/fleet/* → HTTP 200',
            'Admin can access all endpoints → HTTP 200',
            'Dealer cannot access retrain endpoint → HTTP 403',
            'Dealer cannot access model health → HTTP 403',
            'OEM cannot create admin users → HTTP 403',
            '[VULN] Forged OEM token grants OEM access → HTTP 200 (VLN-001)',
            'Invalid role claim in JWT → HTTP 403',
        ]),
        ('D', 'Horizontal Tenant Isolation', 6, [
            'DL001 dealer cannot see DL002 vehicles → scoped response',
            'DL001 dealer cannot trigger alerts for DL002 → scoped response',
            '[VULN] Forged dealer_code=DL002 in token → grants DL002 data (VLN-005)',
            'dealer_code omitted from JWT → HTTP 403',
            'dealer_code empty string → HTTP 403',
            'dealer_code SQL injection pattern → sanitised or 400',
        ]),
        ('E', 'Input Validation', 8, [
            'SQL injection in username field → sanitised',
            'XSS payload in notes field → escaped in response',
            'Oversized request body (>1MB) → HTTP 413 or 400',
            'Missing required fields → HTTP 422',
            'Invalid date format in query params → HTTP 422',
            'Negative values for pagination → HTTP 422',
            'Unicode control characters in text fields → handled',
            'Path traversal in VIN parameter → sanitised',
        ]),
        ('F', 'Information Disclosure', 7, [
            'Error response does not leak stack trace',
            'Error response does not leak internal file paths',
            'Error response does not leak database query',
            '[INFO] /docs endpoint accessibility (informational)',
            '[INFO] /redoc endpoint accessibility (informational)',
            'Server response headers do not expose version',
            '404 response does not distinguish user-not-found from wrong-password',
        ]),
        ('G', 'Transport Security', 3, [
            '[INFO] HTTP vs HTTPS deployment (informational)',
            'Sensitive data not in URL query parameters',
            'Auth token not in server logs',
        ]),
        ('H', 'Auth Hardening', 7, [
            'Rate limiting fires after 5 rapid auth requests → HTTP 429',
            'JWT exp is set (not indefinite)',
            '[VULN] Plaintext password stored → confirmed (VLN-002)',
            'Password not returned in any API response',
            'Login endpoint does not enumerate users (same error for bad user/bad pass)',
            '[INFO] Refresh token mechanism absence (informational)',
            'CORS preflight response → correct Allow-Origin header (not wildcard in prod)',
        ]),
    ]
    add_table(doc,
        ['Section', 'Domain', 'Tests', 'Description'],
        [(s, name, cnt, f'See detailed test list below')
         for s, name, cnt, _ in sections])

    heading(doc, '4. Remediation Priority', level=1)
    add_table(doc,
        ['Priority', 'Vulnerability', 'Effort', 'Risk if Deferred'],
        [
            ('1 (Immediate)', 'VLN-001 Default SECRET_KEY', '30 mins', 'Complete auth system bypass — CRITICAL'),
            ('2 (This Sprint)', 'VLN-002 Plaintext passwords', '2 hours', 'All credentials exposed on breach — CRITICAL'),
            ('3 (This Sprint)', 'VLN-003 CORS wildcard', '15 mins', 'Cross-origin request forgery — HIGH'),
            ('4 (Next Sprint)', 'VLN-004 JWT sub not validated', '1 hour', 'Deleted user tokens remain valid — HIGH'),
            ('5 (Next Sprint)', 'VLN-005 Tenant escalation', '2 hours (depends on VLN-001)', 'Data breach across tenants — HIGH'),
        ])

    doc.save(DOCS / '11_Security_Test_Report.docx')
    print('  ✓ 11_Security_Test_Report.docx')


# ── Doc 12 — Automation Test ──────────────────────────────────────────────────
def doc_automation():
    doc = new_doc()
    cover_page(doc, 'Automation Test Document',
               'Test automation strategy, framework design, patterns, and coverage inventory')

    heading(doc, '1. Automation Strategy')
    body(doc, 'AutoPredict uses a three-tier automation strategy: Playwright E2E for browser '
         'and API testing, axe-core for automated accessibility, and direct HTTP requests for '
         'security testing. No test mocking at the API level — all tests hit the real FastAPI '
         'backend with real CSV data.')

    body(doc, 'Key principle: Tests that exercise the full stack (React → FastAPI → CSV → '
         'LightGBM) find integration bugs that no unit test can catch. The 55–65 second cold '
         'CSV read is accepted as the cost of real integration testing.')

    heading(doc, '2. Test Pyramid', level=1)
    body(doc, 'The pyramid is intentionally inverted toward integration and E2E:')
    pyramid_data = [
        ('Security Tests', 59, RED, 'HTTP attack vector tests (no browser)'),
        ('API Contract Tests', 15, ORG, 'Direct HTTP, full stack, no UI'),
        ('Accessibility Tests', 8, BLUE, 'axe-core WCAG 2.1 AA scans'),
        ('UI E2E Tests', 89, GREEN, 'Browser tests: 11 pages, 5 modules'),
        ('Unit Tests', 3, GRAY, 'Python ML model smoke tests'),
    ]
    add_table(doc,
        ['Test Type', 'Count', 'Tool', 'Scope'],
        [
            (name, str(cnt), 'Playwright + axe-core / Playwright / axe-core / Playwright / pytest', scope)
            for name, cnt, _, scope in pyramid_data
        ])

    heading(doc, '3. Framework Design', level=1)
    heading(doc, '3.1 Directory Structure', level=2)
    structure = [
        'dealer_portal/e2e/',
        '  tests/           — 9 spec files (12-security.spec.ts has 59 tests)',
        '  pages/           — 8 Page Object Models',
        '  fixtures/        — auth.ts: role-based page fixtures',
        '  .auth/           — storageState JSON files (gitignored)',
        '  global-setup.ts  — generates auth state before suite',
        '  global-teardown.ts — cleans up after suite',
        'dealer_portal/playwright.config.ts — config with workers:1',
    ]
    for line in structure:
        p = doc.add_paragraph(line, style='No Spacing')
        for run in p.runs:
            run.font.name = 'Courier New'; run.font.size = Pt(9)
            p.paragraph_format.left_indent = Cm(0.5)

    heading(doc, '3.2 Page Object Model Pattern', level=2)
    body(doc, 'Every page has a corresponding POM class. Tests import the POM — never raw '
         'Playwright locators directly. Each POM provides:')
    bullet(doc, 'Readonly locator properties (evaluated lazily in constructor)')
    bullet(doc, 'goto() — standard navigation method')
    bullet(doc, 'waitFor*() — encapsulates timing logic for that page\'s data loading behaviour')
    bullet(doc, 'Action methods (e.g. setSeverity(), selectModel()) — hide implementation details')

    heading(doc, '3.3 Auth Fixture System', level=2)
    body(doc, '4 storageState files are generated once by global-setup.ts before any test runs. '
         'The auth fixture (e2e/fixtures/auth.ts) provides typed page fixtures:')
    bullet(doc, 'dealerPage — authenticated as DL001 dealer')
    bullet(doc, 'dealer2Page — authenticated as DL002 dealer (for isolation tests)')
    bullet(doc, 'oemPage — authenticated as OEM analyst')
    bullet(doc, 'adminPage — authenticated as system administrator')

    heading(doc, '4. Automation Patterns', level=1)
    patterns = [
        ('Abort-first for static UI tests',
         'Tests that only check headings/buttons abort DATA_API before navigation. '
         'Prevents slow CSV backend requests when data is not needed. '
         'Reduces test time from 65s to ~2s for non-data tests.'),
        ('Load-then-verify for data tests',
         'Tests needing real data skip route.abort() and use waitForData() with 80s timeout. '
         'Promise.any() handles pages that can show either data table or empty state.'),
        ('beforeAll for shared token setup',
         'Security and contract tests obtain auth tokens once in test.beforeAll() '
         'and reuse them across all tests in the describe block.'),
        ('Lenient assertions for informational findings',
         'Security [INFO] tests use expect([200, 404]).toContain(status) — '
         'accepts both outcomes, logs a warning, never fails the suite.'),
        ('[VULN] tests confirm vulnerability',
         'Tests prefixed [VULN] assert HTTP 200 where a secure system would return 403/401. '
         'These tests pass to confirm the vulnerability exists.'),
    ]
    add_table(doc,
        ['Pattern', 'Description'],
        patterns,
        col_widths=[Inches(2.2), Inches(4.3)])

    heading(doc, '5. Anti-Patterns Avoided', level=1)
    anti = [
        ("Glob '**/api/**' for route abort",
         'Matches Vite source file at /src/api/client.ts — prevents React mounting. Use regex instead.'),
        ('Hardcoded sleep() for waits',
         'Flaky and slow. Use waitFor({ state: "visible" }) with explicit timeout.'),
        ('Shared browser state between tests',
         'Test pollution. Each test gets fresh page from role fixture.'),
        ('Login steps in every test',
         'Overhead and auth test failure cascades. Global setup generates storageState once.'),
        ('Parallel workers with CSV backend',
         'Causes backend queue buildup and 25–45s timeouts. workers:1 always.'),
        ('actionTimeout in playwright.config.ts',
         'Propagates to APIRequestContext — security tests timeout on slow endpoints. Omit it.'),
    ]
    add_table(doc,
        ['Anti-Pattern', 'Why Avoided'],
        anti,
        col_widths=[Inches(2.4), Inches(4.1)])

    heading(doc, '6. CI Integration', level=1)
    add_table(doc,
        ['Step', 'Command', 'Purpose'],
        [
            ('Backend start', 'uvicorn api.main:app --port 8001 &', 'Start FastAPI server'),
            ('Frontend build', 'npm run build && npm run preview &', 'Serve production bundle'),
            ('Browser install', 'npx playwright install chromium --with-deps', 'Install Chromium once'),
            ('Run suite', 'npx playwright test', 'Execute all 174 tests'),
            ('Upload report', 'Upload playwright-report/ artifact', 'Post-run HTML report'),
        ])

    heading(doc, '7. Coverage Gaps', level=1)
    body(doc, 'Three test files remain to be written:')
    add_table(doc,
        ['Missing File', 'What to Test', 'Estimated Tests'],
        [
            ('04-service-bay.spec.ts', 'Bay grid, booking modal, appointment list', '10'),
            ('05-inventory.spec.ts', 'All 7 tabs, demand forecast, reorder plan', '12'),
            ('07-oem-models.spec.ts', 'Model cards, EDA sub-tabs, What-If form', '8'),
        ])

    doc.save(DOCS / '12_Automation_Test_Document.docx')
    print('  ✓ 12_Automation_Test_Document.docx')


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    print('Generating Word documents...')
    doc_tech_arch()
    doc_product_spec()
    doc_prd()
    doc_user_journey()
    doc_roadmap()
    doc_devops()
    doc_frontend()
    doc_backend()
    doc_test_cases()
    doc_playwright()
    doc_security()
    doc_automation()
    print(f'\nAll documents saved to: {DOCS}')
